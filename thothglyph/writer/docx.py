from __future__ import annotations
from typing import Dict, List, Optional
import importlib
import os
import tempfile
import cairosvg
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph
# from docx.enum.text import WD_TAB_ALIGNMENT
# from docx.enum.text import WD_TAB_LEADER
from thothglyph.error import ThothglyphError
from thothglyph.writer.writer import Writer
from thothglyph.node import nd
from thothglyph.node import logging

logger = logging.getLogger(__file__)

# monkey patch
if True:
    from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne
    # from docx.oxml.simpletypes import XsdInt
    from docx.oxml import register_element_cls

    from types import MethodType
    # from docx.parts.numbering import _NumberingDefinitions
    from docx.oxml.numbering import CT_Num
    # from docx.oxml.numbering import CT_Numbering

    class CT_TrPr(BaseOxmlElement):
        tblHeader = ZeroOrOne('w:tblHeader')
    register_element_cls('w:trPr', CT_TrPr)


class DocxWriter(Writer):
    target = 'docx'
    ext = 'docx'

    stylename: Dict[str, str] = {
        # paragraph styles
        'title': 'Title',
        'subtitle': 'Subtitle',
        'heading_1': 'Heading 1',
        'heading_2': 'Heading 2',
        'heading_3': 'Heading 3',
        'heading_4': 'Heading 4',
        'heading_5': 'Heading 5',
        'heading_6': 'Heading 6',
        'heading_7': 'Heading 7',
        'bullet_list': 'Bullet List',
        'enumerated_list': 'Enumerated List',
        '_empty_bullet_list': 'Empty Bullet List',
        'definition_list': 'Definition List',
        'definition_list_item': 'Definition List Item',
        'field_list': 'Field List',
        'field_list_item': 'Field List Item',
        'option_list': 'Option List',
        'option_list_item': 'Option List Item',
        'literal_block': 'Literal Block',
        'code_block': 'Literal Block',
        'block_quote': 'Quote',
        'line_block': 'Line Block',
        'doctest_block': 'Doctest Block',
        'transition': 'Horizontal Line',
        'table_caption': 'Caption',
        'image_caption': 'Caption',
        'code_block_caption': 'Code Block Caption',
        # character styles
        'strong': 'Strong',
        'emphasis': 'Emphasis',
        'literal_emphasis': 'Literal Emphasis',
        'subscript': 'Subscript',
        'superscript': 'Superscript',
        'title_reference': 'Book Title',
        'literal': 'Literal',
        'code': 'Literal',
        'reference': 'Hyperlink',
        'footnote_reference': 'Default Paragraph Font',
        # table styles
        'table_normal': 'Sphinx Table Normal',
        'table_list': 'Sphinx Table List',
    }

    decoration_table: Dict[str, str] = {
        'EMPHASIS': 'emphasis',
        'STRONG': 'strong',
        'MARKED': 'strong',
        'STRIKE': 'strong',
        'VAR': 'strong',
        'CODE': 'code',
        'SUP': 'superscript',
        'SUB': 'subscript',
    }

    def __init__(self):
        super().__init__()
        self.tmpdirname: Optional[str] = None

        self.numbered = 0
        self.numbered_level = 0
        self.section_level = 0
        self.section_numIds = list()
        self.initial_header_level = 0  # int(self.settings.initial_header_level)
        # docx paragraph properties
        self.p: Optional[Paragraph] = None
        self.p_parents = list()
        self.p_style: List[str] = list()
        self.p_level = 0
        self.numIds = list()
        self.is_first_list_item = False
        # special paragraphs
        self.tables = list()
        self.item_width_rate = 0.8
        # docx run properties
        self.r = None
        self.r_style: Optional[str] = None

    def parse(self, node: nd.ASTNode) -> None:
        template_dir = self.template_dir()
        target = self.target
        theme = self.theme()
        template_path = os.path.join(template_dir, target, theme, 'style.docx')
        if not os.path.exists(template_path):
            raise ThothglyphError('template not found: {}'.format(template_path))
        # t = template.replace('{', '{{').replace('}', '}}')
        # t = re.sub(r'\$\{\{([^}]+)\}\}', r'{\1}', t)
        # self.data = t.format(doc=self.template_docdata)
        assert self.tmpdirname
        for checkbox in ('check_en', 'check_im', 'check_dis'):
            cairosvg.svg2png(
                url=os.path.join(template_dir, 'common', f'{checkbox}.svg'),
                write_to=os.path.join(self.tmpdirname, f'{checkbox}.png'),
                scale=0.625,
            )
        self.data: Document = Document(template_path)
        self.data._body.clear_content()
        self.p_parents.append(self.data)
        self.section_numIds = [self._get_new_num(abstractNumId=12)]
        super().parse(node)

    def write(self, fpath: str, node: nd.ASTNode) -> None:
        clsname = self.__class__.__name__
        logger.info('{}: write document'.format(clsname))
        self.rootnode = node
        with tempfile.TemporaryDirectory() as tmpdirname:
            self.tmpdirname = tmpdirname
            self.parse(node)
            self.data.save(fpath)
        self.tmpdirname = None

    def _add_paragraph(
        self,
        text: Optional[str] = None,
        style: Optional[str] = None
    ) -> Optional[Paragraph]:
        p = None
        try:
            if isinstance(style, list):
                p = self.p_parents[-1].add_paragraph(text, style[-1])
            else:
                p = self.p_parents[-1].add_paragraph(text, style)
        except Exception:
            p = self.p_parents[-1].add_paragraph(text, 'Normal')
        if self.p_level > 0:
            self._multilevel_list_numbering(p, self.p_level - 1, 15)
        return p

    def _add_run(self, text=None, style=None):
        r = None
        if self.p:
            try:
                r = self.p.add_run(text, style)
            except Exception:
                r = self.p.add_run(text, 'Default Paragraph Font')
        return r

    def _get_new_num(self, abstractNumId):

        def add_num(self, abstractNum_id, restart=False):
            next_num_id = self._next_numId
            num = CT_Num.new(next_num_id, abstractNum_id)
            if restart:
                num.add_lvlOverride(ilvl=0).add_startOverride(1)
            return self._insert_num(num)
        numbering = self.data._part.numbering_part.numbering_definitions._numbering
        numbering.add_num = MethodType(add_num, numbering)
        num = numbering.add_num(abstractNumId, True).numId
        return num

    def _multilevel_list_numbering(self, paragraph, ilvl, numId):
        # monkey patch
        pfmt = paragraph.paragraph_format
        numPr = pfmt._element.get_or_add_pPr().get_or_add_numPr()
        numPr.get_or_add_ilvl().val = ilvl
        numPr.get_or_add_numId().val = numId

    def visit_section(self, node):
        # _id = node.id or node.title.replace(' ', '_')
        if node.opts.get('nonum'):
            title = node.title
        else:
            title = '{}. {}'.format(node.sectnum, node.title)
        hlevel = node.level if node.level <= 6 else 6
        if hlevel <= 1:
            lastp = self.data.paragraphs[-1] if len(self.data.paragraphs) > 0 else None
            if lastp:
                lastp.add_run().add_break(WD_BREAK.PAGE)
        self.data.add_heading(title, level=hlevel)

    def leave_section(self, node):
        self.p = None
        self.r = None

    def visit_bulletlistblock(self, node):
        self.p_level += 1
        numId = self._get_new_num(abstractNumId=11)
        self.numIds.append(numId)

    def leave_bulletlistblock(self, node):
        self.p_level -= 1
        self.numIds.pop()

    def visit_orderedlistblock(self, node):
        self.p_level += 1
        numId = self._get_new_num(abstractNumId=15)
        self.numIds.append(numId)

    def leave_orderedlistblock(self, node):
        self.p_level -= 1
        self.numIds.pop()

    def visit_descriptionlistblock(self, node):
        pass

    def leave_descriptionlistblock(self, node):
        pass

    def visit_checklistblock(self, node):
        pass

    def leave_checklistblock(self, node):
        pass

    def visit_footnotelistblock(self, node):
        pass

    def leave_footnotelistblock(self, node):
        pass

    def visit_referencelistblock(self, node):
        pass

    def leave_referencelistblock(self, node):
        pass

    def visit_listitem(self, node):
        item = node
        if isinstance(item.parent, nd.FootnoteListBlockNode):
            r_style = self.stylename['footnote_reference']
            # url = '#fn.{}-{}'.format(item.treeindex()[1], item.fn_num)
            text = item.fn_num  # item.title
            self.p = self._add_paragraph()
            self.r = self._add_run('[{}] '.format(text), r_style)
        elif isinstance(item.parent, nd.ReferenceListBlockNode):
            r_style = self.stylename['reference']
            # url = '#ref.{}'.format(item.ref_num)
            text = item.ref_num  # item.title
            self.p = self._add_paragraph()
            self.r = self._add_run('[{}] '.format(text), r_style)

    def leave_listitem(self, node):
        self.p = None
        self.r = None

    def visit_quoteblock(self, node):
        self.p_style.append(self.stylename['block_quote'])
        self.p_level += 1

    def leave_quoteblock(self, node):
        self.p_style.pop()
        self.p_level -= 1

    def visit_codeblock(self, node: nd.ASTNode):
        self.r_style = self.stylename['code_block']
        self.p = self._add_paragraph('', self.r_style)
        self.r = None

    def leave_codeblock(self, node):
        self.p = None
        self.r_style = None

    def visit_figureblock(self, node):
        # if node.align == 'l':
        #     style = 'style="text-align:left;"'
        # elif node.align == 'c':
        #     style = 'style="text-align:center;"'
        # elif node.align == 'r':
        #     style = 'style="text-align:right;"'
        if isinstance(node.children[0], nd.TableBlockNode):
            self.r_style = self.stylename['table_caption']
            text = '{} {}'.format(node.fignum, node.caption)
            self.p = self._add_paragraph(text, self.r_style)
            self.p = None
            self.r = None
        else:
            pass

    def leave_figureblock(self, node):
        if isinstance(node.children[0], nd.TableBlockNode):
            pass
        else:
            self.r_style = self.stylename['image_caption']
            text = '{} {}'.format(node.fignum, node.caption)
            self.p = self._add_paragraph(text, self.r_style)
            self.p = None
            self.r = None

    def _add_paragraph_between_table(self, node):
        index = node.parent.children.index(node)
        prev_node = node.parent.children[index - 1]
        if (isinstance(prev_node, nd.TableBlockNode)):
            # or isinstance(prev_node, nodes.field_list)
            # or isinstance(prev_node, nodes.option_list)):
            self.data.add_paragraph('')

    def visit_tableblock(self, node):
        self._add_paragraph_between_table(node)
        table = self.p_parents[-1].add_table(rows=node.row, cols=node.col)
        align = node.align
        if isinstance(node.parent, nd.FigureBlockNode):
            align = node.parent.align
        if align:
            if align == 'l':
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
            elif align == 'c':
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            elif align == 'r':
                table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            else:
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for raw in table.rows:
            for c in raw._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0
        if node.headers == 0:
            table.style = self.stylename['table_normal']
        else:
            table.style = self.stylename['table_list']
            for i in range(node.headers):
                trPr = table.rows[i]._tr.get_or_add_trPr()
                trPr.get_or_add_tblHeader()
        self.tables.append(table)

    def leave_tableblock(self, node):
        self.tables.pop()

    def visit_tablerow(self, node):
        pass

    def leave_tablerow(self, node):
        pass

    def visit_tablecell(self, node):
        table = self.tables[-1]
        r, c = node.parent.idx, node.idx
        cell = table.cell(r, c)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]
        s = node.size
        if s.x > 1 or s.y > 1:
            cell.merge(table.cell(r + s.y - 1, c + s.x - 1))

    def leave_tablecell(self, node):
        self.p_parents.pop()

    def visit_customblock(self, node):
        if node.ext == '':
            self.r_style = self.stylename['code_block']
            self.p = self._add_paragraph(node.text, self.r_style)
            self.r = None
        else:
            try:
                extpath = 'thothglyph.ext.{}'.format(node.ext)
                extmodule = importlib.import_module(extpath)
                extmodule.customblock_write_docx(self, node)
            except Exception:
                self.r_style = self.stylename['code_block']
                self.p = self._add_paragraph(node.text, self.r_style)
                self.r = None

    def leave_customblock(self, node):
        pass

    def visit_horizonblock(self, node):
        self.data.add_paragraph('', style=self.stylename['transition'])
        self.p = None

    def leave_horizonblock(self, node):
        pass

    def visit_paragraph(self, node):
        if not self.p:
            self.p = self._add_paragraph()
        if isinstance(node.parent, nd.ListItemNode):
            item = node.parent
            is_first_paragraph = item.children.index(node) == 0
            if is_first_paragraph and self.numIds:
                self._multilevel_list_numbering(self.p, item.level - 1, self.numIds[-1])
            else:
                self._multilevel_list_numbering(self.p, item.level - 1, 15)
            self.is_first_list_item = False
            if isinstance(item.parent, nd.CheckListBlockNode):
                r_style = self.stylename['strong']
                assert self.tmpdirname
                if item.marker == 'x':
                    checkbox = os.path.join(self.tmpdirname, 'check_en.png')
                    # self.r = self._add_run('[v] ', r_style)
                elif item.marker == '-':
                    checkbox = os.path.join(self.tmpdirname, 'check_im.png')
                    # self.r = self._add_run('[-] ', r_style)
                else:
                    checkbox = os.path.join(self.tmpdirname, 'check_dis.png')
                    # self.r = self._add_run('[ ] ', r_style)
                self.r = self._add_run('')
                self.r.add_picture(checkbox)
                self.r = self._add_run(' ')
        self.r = self.p.add_run()

    def leave_paragraph(self, node):
        self.p = None
        self.r = None

    def visit_title(self, node):
        if not self.p:
            self.p = self._add_paragraph()
        if isinstance(node.parent, nd.DescriptionListBlockNode):
            self.r_style = self.stylename['strong']

    def leave_title(self, node):
        self.r_style = None
        self.r = self._add_run(' ')

    def visit_decorationrole(self, node):
        stylekey = self.decoration_table[node.role]
        self.r_style = self.stylename[stylekey]

    def visit_role(self, node):
        if node.role == '':
            self.r_style = self.stylename['code_block']
            self.r = self._add_run(node.value, self.r_style)
            self.r = None
        else:
            try:
                extpath = 'thothglyph.ext.{}'.format(node.role)
                extmodule = importlib.import_module(extpath)
                extmodule.role_write_docx(self, node)
            except Exception:
                self.r_style = self.stylename['code_block']
                self.r = self._add_run(node.value, self.r_style)
                self.r = None

    def leave_role(self, node):
        pass

    def visit_imagerole(self, node):
        optdict: Dict[str, str] = dict()
        if 'w' in node.opts:
            optdict['width'] = node.opts['w']
        optlist = ['{}="{}"'.format(k, v) for k, v in optdict.items()]
        options = ' '.join(optlist)

        image_fullpath = os.path.abspath(node.value)
        imagedir, imagefname = os.path.splitext(image_fullpath)
        fname, ext = os.path.splitext(image_fullpath)
        if ext == '.svg':
            assert self.tmpdirname
            tmpimage_path = os.path.join(self.tmpdirname, f'{fname}.png')
            cairosvg.svg2png(
                url=image_fullpath,
                write_to=tmpimage_path,
                scale=0.625,
            )
            image_fullpath = tmpimage_path
        if isinstance(node.parent, nd.ParagraphNode):
            if self.r:
                self.r.add_picture(image_fullpath)
        elif isinstance(node.parent, nd.FigureBlockNode):
            if self.r:
                self.r.add_picture(image_fullpath)
        else:
            p = self._add_paragraph()
            if p:
                r = p.add_run()
                r.add_picture(image_fullpath)

    def leave_imagerole(self, node):
        pass

    def visit_kbdrole(self, node):
        # value = ' + '.join(['<kbd>{}</kbd>'.format(v) for v in node.value])
        r_style = self.stylename['code']
        self._add_run(node.value, r_style)

    def leave_kbdrole(self, node):
        pass

    def visit_btnrole(self, node):
        # value = '<kbd>{}</kbd>'.format(node.value)
        r_style = self.stylename['code']
        self._add_run(node.value, r_style)

    def leave_btnrole(self, node):
        pass

    def visit_menurole(self, node):
        # value = ' > '.join(['<span class="menu">{}</span>'.format(v) for v in node.value])
        r_style = self.stylename['code']
        self._add_run(node.value, r_style)

    def leave_menurole(self, node):
        pass

    def visit_link(self, node):
        # if '://' in node.value:
        #     blank = 'target=”_blank”'
        #     url = node.value
        #     text = node.opts[0] if node.opts[0] else node.value
        # else:
        #     blank = ''
        #     url = '#' + node.target_id
        #     text = node.opts[0] if node.opts[0] else node.target_title
        r_style = self.stylename['reference']
        self._add_run(node.value, r_style)

    def leave_link(self, node):
        pass

    def visit_text(self, node):
        if self.p:
            self.r = self._add_run(node.text, self.r_style)
        self.r_style = None

    def visit_other(self, node):
        pass

    def leave_other(self, node):
        pass
